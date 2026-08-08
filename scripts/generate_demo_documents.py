from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
CORPUS = ROOT / "examples" / "acme-corp"

NAVY = "17324D"
TEAL = "149A9A"
PALE = "EAF3F3"
INK = "263746"
MUTED = "607080"


def _shade(cell: object, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_margins(cell: object, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc  # type: ignore[attr-defined]
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def create_docx(
    filename: str,
    number: str,
    subject: str,
    effective: str,
    scope: str,
    priority: str,
    clauses: list[tuple[str, str]],
) -> None:
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, color in (("Title", 22, NAVY), ("Heading 1", 13, NAVY)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "ACME CORP  /  PEOPLE OPERATIONS"
    header.style = styles["Normal"]
    header.runs[0].font.bold = True
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Внутренний документ • контролируемая копия")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)

    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(4)
    run = eyebrow.add_run(f"ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ  ·  {number}")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    title = document.add_paragraph(subject, style="Title")
    title.paragraph_format.space_after = Pt(12)

    intro = document.add_paragraph(
        "Настоящее дополнительное соглашение устанавливает специальный порядок и применяется "
        "вместе с базовым корпоративным документом."
    )
    intro.paragraph_format.space_after = Pt(10)

    table = document.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(12.4)
    for row, (label, value) in zip(
        table.rows,
        (("Дата вступления", effective), ("Область действия", scope), ("Приоритет", priority)),
        strict=True,
    ):
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = Cm(12.4)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
        _shade(row.cells[0], NAVY)
        _shade(row.cells[1], PALE)
        left_run = row.cells[0].paragraphs[0].add_run(label.upper())
        left_run.font.name = "Arial"
        left_run.font.size = Pt(8)
        left_run.font.bold = True
        left_run.font.color.rgb = RGBColor(255, 255, 255)
        right_run = row.cells[1].paragraphs[0].add_run(value)
        right_run.font.name = "Arial"
        right_run.font.size = Pt(9)
        right_run.font.color.rgb = RGBColor.from_string(INK)

    for heading, body in clauses:
        document.add_paragraph(heading, style="Heading 1")
        document.add_paragraph(body)

    signature = document.add_table(rows=1, cols=2)
    signature.autofit = False
    signature.columns[0].width = Cm(8.3)
    signature.columns[1].width = Cm(8.3)
    for cell, text in zip(signature.rows[0].cells, ("Работодатель: Acme Corp", "Сотрудник / представитель подразделения"), strict=True):
        cell.width = Cm(8.3)
        _set_cell_margins(cell, top=160, bottom=60)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(10)
        run = paragraph.add_run(text + "\n\n________________________")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    document.save(CORPUS / filename)


def _pdf_fonts() -> tuple[str, str]:
    regular_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    pdfmetrics.registerFont(TTFont("AcmeArial", regular_path))
    pdfmetrics.registerFont(TTFont("AcmeArialBold", bold_path))
    return "AcmeArial", "AcmeArialBold"


def create_pdf(
    filename: str,
    number: str,
    subject: str,
    effective: str,
    scope: str,
    priority: str,
    clauses: list[tuple[str, str]],
) -> None:
    regular, bold = _pdf_fonts()
    output = CORPUS / filename
    document = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=subject,
        author="Acme Corp",
    )
    styles = getSampleStyleSheet()
    eyebrow = ParagraphStyle(
        "Eyebrow", parent=styles["Normal"], fontName=bold, fontSize=8, textColor=colors.HexColor("#149A9A"), leading=10
    )
    title = ParagraphStyle(
        "Title", parent=styles["Title"], fontName=bold, fontSize=21, leading=24, alignment=TA_LEFT,
        textColor=colors.HexColor("#17324D"), spaceAfter=10
    )
    heading = ParagraphStyle(
        "Heading", parent=styles["Heading2"], fontName=bold, fontSize=12, leading=14,
        textColor=colors.HexColor("#17324D"), spaceBefore=9, spaceAfter=4
    )
    body = ParagraphStyle(
        "Body", parent=styles["BodyText"], fontName=regular, fontSize=9.5, leading=13,
        textColor=colors.HexColor("#263746"), spaceAfter=5
    )
    meta_label = ParagraphStyle("MetaLabel", parent=body, fontName=bold, fontSize=7.5, textColor=colors.white, leading=9)
    meta_value = ParagraphStyle("MetaValue", parent=body, fontSize=8.5, leading=11)

    story: list[object] = [
        Paragraph("ACME CORP / PEOPLE OPERATIONS", eyebrow),
        Spacer(1, 4 * mm),
        Paragraph(f"ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ · {number}", eyebrow),
        Paragraph(subject, title),
        Paragraph(
            "Настоящее дополнительное соглашение устанавливает специальный порядок и применяется вместе с базовым корпоративным документом.",
            body,
        ),
        Spacer(1, 3 * mm),
    ]
    rows = [
        [Paragraph("ДАТА ВСТУПЛЕНИЯ", meta_label), Paragraph(effective, meta_value)],
        [Paragraph("ОБЛАСТЬ ДЕЙСТВИЯ", meta_label), Paragraph(scope, meta_value)],
        [Paragraph("ПРИОРИТЕТ", meta_label), Paragraph(priority, meta_value)],
    ]
    metadata = Table(rows, colWidths=[43 * mm, 117 * mm], hAlign="LEFT")
    metadata.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#17324D")),
                ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#EAF3F3")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LINEBELOW", (0, 0), (-1, -2), 0.5, colors.white),
            ]
        )
    )
    story.extend([metadata, Spacer(1, 3 * mm)])
    for clause_heading, clause_body in clauses:
        story.extend([Paragraph(clause_heading, heading), Paragraph(clause_body, body)])
    story.extend(
        [
            Spacer(1, 8 * mm),
            Paragraph("Работодатель: Acme Corp  ____________________", body),
            Paragraph("Сотрудник / представитель подразделения  ____________________", body),
        ]
    )
    document.build(story)


def create_reference_docx(
    filename: str,
    kicker: str,
    title_text: str,
    subtitle: str,
    metadata: list[tuple[str, str]],
    sections: list[tuple[str, list[str]]],
) -> None:
    """Create a polished Acme reference document using the repository visual system."""
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, before, after in (
        ("Heading 1", 15, 14, 7),
        ("Heading 2", 12, 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header_run = header.add_run("ACME CORP  /  PEOPLE OPERATIONS")
    header_run.font.name = "Arial"
    header_run.font.size = Pt(8)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Учебный документ • синтетические данные")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)

    kicker_paragraph = document.add_paragraph()
    kicker_run = kicker_paragraph.add_run(kicker.upper())
    kicker_run.font.name = "Arial"
    kicker_run.font.size = Pt(9)
    kicker_run.font.bold = True
    kicker_run.font.color.rgb = RGBColor.from_string(TEAL)
    kicker_paragraph.paragraph_format.space_after = Pt(5)

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title_text)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    title_paragraph.paragraph_format.space_after = Pt(5)

    subtitle_paragraph = document.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle_paragraph.paragraph_format.space_after = Pt(12)

    table = document.add_table(rows=len(metadata), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(12.0)
    for row, (label, value) in zip(table.rows, metadata, strict=True):
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = Cm(12.0)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=100, bottom=100)
        _shade(row.cells[0], NAVY)
        _shade(row.cells[1], PALE)
        label_run = row.cells[0].paragraphs[0].add_run(label.upper())
        label_run.font.name = "Arial"
        label_run.font.size = Pt(8)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(255, 255, 255)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        value_run.font.name = "Arial"
        value_run.font.size = Pt(9)
        value_run.font.color.rgb = RGBColor.from_string(INK)

    for heading, paragraphs in sections:
        document.add_paragraph(heading, style="Heading 1")
        for paragraph_text in paragraphs:
            document.add_paragraph(paragraph_text)

    document.save(CORPUS / filename)


def create_reference_pdf(
    filename: str,
    kicker: str,
    title_text: str,
    subtitle: str,
    sections: list[tuple[str, list[str]]],
) -> None:
    regular, bold = _pdf_fonts()
    output = CORPUS / filename
    document = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title_text,
        author="Acme Corp",
    )
    styles = getSampleStyleSheet()
    eyebrow = ParagraphStyle(
        "ReferenceEyebrow",
        parent=styles["Normal"],
        fontName=bold,
        fontSize=8,
        textColor=colors.HexColor("#149A9A"),
        leading=10,
        spaceAfter=4,
    )
    title = ParagraphStyle(
        "ReferenceTitle",
        parent=styles["Title"],
        fontName=bold,
        fontSize=23,
        leading=27,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#17324D"),
        spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "ReferenceSubtitle",
        parent=styles["BodyText"],
        fontName=regular,
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#607080"),
        spaceAfter=10,
    )
    heading = ParagraphStyle(
        "ReferenceHeading",
        parent=styles["Heading2"],
        fontName=bold,
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#17324D"),
        spaceBefore=10,
        spaceAfter=5,
    )
    body = ParagraphStyle(
        "ReferenceBody",
        parent=styles["BodyText"],
        fontName=regular,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#263746"),
        spaceAfter=6,
    )
    story: list[object] = [
        Paragraph("ACME CORP / ВНУТРЕННИЙ ДОКУМЕНТ", eyebrow),
        Spacer(1, 4 * mm),
        Paragraph(kicker.upper(), eyebrow),
        Paragraph(title_text, title),
        Paragraph(subtitle, subtitle_style),
    ]
    for section_heading, paragraphs in sections:
        story.append(Paragraph(section_heading, heading))
        story.extend(Paragraph(text, body) for text in paragraphs)
    story.extend(
        [
            Spacer(1, 6 * mm),
            Paragraph("Учебный документ • синтетические данные • редакция 01.10.2025", eyebrow),
        ]
    )
    document.build(story)


def main() -> None:
    create_reference_docx(
        "employment_contract_template.docx",
        "Трудовой договор · учебный образец",
        "Трудовой договор сотрудника Acme Corp",
        "Шаблон для демонстрации поиска условий труда; не является документом для подписания.",
        [
            ("Работодатель", "Acme Corp, демонстрационная организация"),
            ("Сотрудник", "Учебный пример без персональных данных"),
            ("Срок", "Бессрочный договор; начало работы — дата в кадровом приказе"),
            ("Испытание", "Три месяца, если иное не указано в индивидуальном соглашении"),
        ],
        [
            (
                "1. Трудовая функция",
                [
                    "Сотрудник выполняет обязанности по должностной инструкции своего подразделения, "
                    "соблюдает корпоративные политики и взаимодействует с назначенным руководителем.",
                ],
            ),
            (
                "2. Рабочее время и формат",
                [
                    "Нормальная продолжительность — 40 часов в неделю. Базовый график: понедельник–пятница, "
                    "09:00–18:00, перерыв 13:00–14:00. Сменный или гибридный режим определяется политикой "
                    "форматов работы и действующим дополнительным соглашением.",
                    "Место работы может быть офисным, удалённым или гибридным. Специальное соглашение по "
                    "подразделению имеет приоритет в пределах своей области действия и срока.",
                ],
            ),
            (
                "3. Условия и обязанности",
                [
                    "Работодатель предоставляет управляемое устройство и необходимые доступы. Сотрудник "
                    "обязан защищать данные, проходить обязательное обучение и сообщать об инцидентах.",
                    "Изменение должности, оплаты, графика или постоянного места работы оформляется письменно.",
                ],
            ),
            (
                "4. Завершение отношений",
                [
                    "В последний рабочий день сотрудник возвращает оборудование, передаёт рабочие материалы "
                    "и подтверждает закрытие доступов через Service Desk.",
                ],
            ),
        ],
    )
    create_reference_docx(
        "workplace_setup_guide.docx",
        "Инструкция · рабочее место",
        "Первичная настройка рабочего компьютера",
        "Пошаговый чек-лист для самостоятельной подготовки управляемого устройства.",
        [
            ("Владелец", "IT Operations"),
            ("Версия", "1.4 от 1 октября 2025 года"),
            ("Применение", "Windows 11 и macOS на оборудовании Acme Corp"),
            ("Поддержка", "Service Desk → категория «Рабочее место»"),
        ],
        [
            (
                "1. До первого входа",
                [
                    "Сверьте инвентарный номер с актом выдачи, подключите питание и доверенную сеть. Не "
                    "используйте личную учётную запись Apple, Microsoft или Google для первичной настройки.",
                ],
            ),
            (
                "2. Управление и защита",
                [
                    "Активируйте ACME ID и MFA, дождитесь регистрации устройства в MDM, включите шифрование "
                    "диска и автоматическую блокировку через 5 минут. Обновления операционной системы должны "
                    "установиться до работы с корпоративными данными.",
                ],
            ),
            (
                "3. Рабочее окружение",
                [
                    "Установите CorpMail, Connect, Service Desk и Corporate VPN из каталога приложений. "
                    "Инженеры дополнительно получают GitLab и инструменты разработки после согласования.",
                    "Проверьте вход через VPN, отправьте тестовое письмо наставнику и создайте подпись по "
                    "корпоративному шаблону.",
                ],
            ),
            (
                "4. Проверка готовности",
                [
                    "Устройство считается готовым, если MDM показывает статус compliant, MFA работает, "
                    "CorpMail синхронизируется, VPN подключается, а требуемые системе доступны без общих паролей.",
                ],
            ),
        ],
    )
    create_reference_pdf(
        "culture_code.pdf",
        "Культурный код и ценности",
        "Как мы работаем в Acme Corp",
        "Миссия, ценности и принципы взаимодействия для ежедневных решений.",
        [
            (
                "Миссия",
                [
                    "Делать сложное корпоративное знание доступным, проверяемым и полезным в момент решения.",
                ],
            ),
            (
                "Ясность",
                [
                    "Формулируем проблему, владельца и критерий результата. Решение фиксируем там, где его "
                    "найдёт следующая команда, а не только участники встречи.",
                ],
            ),
            (
                "Ответственность",
                [
                    "Берём задачу до проверяемого результата, заранее поднимаем риск и не скрываем ошибку. "
                    "Право принять решение сопровождается обязанностью объяснить его основания.",
                ],
            ),
            (
                "Уважение и сотрудничество",
                [
                    "Критикуем идею, а не человека. Договариваемся о канале и сроке ответа, учитываем "
                    "разницу часовых поясов и не превращаем срочность одного отдела в постоянную срочность другого.",
                ],
            ),
            (
                "Доказательность",
                [
                    "Отделяем факт от предположения, указываем источник и проверяем эффект после изменения. "
                    "Если данных недостаточно, задаём уточняющий вопрос вместо уверенного ответа.",
                ],
            ),
        ],
    )
    create_docx(
        "additional_agreement_finance_office.docx",
        "RW-FIN-2025",
        "Офисный режим финансового блока",
        "1 октября 2025 года; действует до 31 декабря 2026 года.",
        "Сотрудники бухгалтерии, казначейства и финансового контроля.",
        "Имеет приоритет над базовой политикой форматов работы в пределах финансового блока.",
        [
            (
                "1. Специальный режим",
                "Понедельник–четверг являются обязательными офисными днями. Пятница может быть удалённым "
                "днём после фиксации графика у руководителя до 16:00 четверга.",
            ),
            (
                "2. Закрытие периода",
                "В последние три рабочих дня месяца сотрудники работают из офиса независимо от обычного "
                "пятничного режима, если финансовый директор письменно не установил исключение.",
            ),
            (
                "3. Защита данных",
                "Платёжные реестры и банковские документы обрабатываются только с управляемого устройства "
                "через VPN; локальное хранение после завершения задачи запрещено.",
            ),
        ],
    )
    create_docx(
        "additional_agreement_support_shifts.docx",
        "SHIFT-SUPPORT-2025",
        "Сменный режим клиентской поддержки",
        "15 октября 2025 года; действует до 31 декабря 2026 года.",
        "Специалисты первой и второй линии клиентской поддержки после завершения адаптации.",
        "Заменяет базовый пятидневный график и правила места работы только для утверждённых смен.",
        [
            (
                "1. График смен",
                "Применяется цикл два рабочих дня через два выходных. Дневная смена: 08:00–20:00; ночная "
                "смена: 20:00–08:00. График публикуется не позднее чем за 14 календарных дней.",
            ),
            (
                "2. Место работы",
                "Ночные смены выполняются удалённо через Corporate VPN. Дневная смена специалиста первой "
                "линии проходит в офисе; специалист второй линии может работать удалённо до двух дневных "
                "смен в месяц после согласования руководителя.",
            ),
            (
                "3. Передача смены",
                "За 20 минут до окончания сотрудник обновляет журнал обращений и передаёт открытые инциденты "
                "следующей смене. Критический инцидент нельзя оставлять без назначенного владельца.",
            ),
        ],
    )
    create_pdf(
        "additional_agreement_design_hybrid.pdf",
        "RW-DESIGN-2025",
        "Гибридный режим дизайн-команды",
        "1 октября 2025 года; действует до 31 декабря 2026 года.",
        "UX/UI-дизайнеры и исследователи продуктового блока.",
        "Имеет приоритет над базовым лимитом удалённой работы; очные исследования сохраняют приоритет.",
        [
            (
                "1. Специальный режим",
                "Допускается до трёх удалённых рабочих дней в неделю. Вторник и четверг являются общими "
                "офисными днями дизайн-команды.",
            ),
            (
                "2. Исследования",
                "Модерируемое исследование в офисной лаборатории может потребовать дополнительного очного "
                "дня с уведомлением не позднее чем за два рабочих дня.",
            ),
            (
                "3. Материалы",
                "Записи исследований и макеты с персональными данными хранятся только в утверждённых "
                "пространствах; перенос в личные облака запрещён.",
            ),
        ],
    )
    create_docx(
        "additional_agreement_remote_sales.docx",
        "RW-SALES-2025",
        "Режим удалённой работы отдела продаж",
        "1 августа 2025 года; действует до 31 декабря 2026 года.",
        "Штатные сотрудники отдела продаж коммерческого блока, включая руководителей групп.",
        "Имеет приоритет над базовым лимитом удалённой работы в пределах области действия и срока.",
        [
            ("1. Специальный режим", "Сотрудник может работать удалённо до трёх рабочих дней в неделю. Вторник и четверг являются обязательными офисными днями для встреч с командой."),
            ("2. Согласование", "График на следующую неделю фиксируется в календаре до 16:00 пятницы. Руководитель вправе назначить дополнительный офисный день для клиентской встречи с уведомлением за два рабочих дня."),
            ("3. Сохраняющиеся требования", "Правила учёта рабочего времени, защиты клиентских данных и доступности в корпоративном мессенджере применяются без изменений."),
        ],
    )
    create_docx(
        "additional_agreement_bonus_sales.docx",
        "BONUS-SALES-Q3-2025",
        "Квартальный бонус отдела продаж",
        "1 июля 2025 года; применяется к результатам III квартала 2025 года.",
        "Штатные сотрудники отдела продаж, имеющие утверждённый индивидуальный план выручки.",
        "Имеет приоритет над базовым целевым бонусом 10% только при расчёте бонуса за III квартал 2025 года.",
        [
            ("1. Целевой бонус", "При выполнении индивидуального плана выручки на 100% целевой бонус составляет 15% квартального оклада."),
            ("2. Повышенный результат", "При выполнении плана свыше 120% бонус составляет 20% квартального оклада. Между 100% и 120% применяется линейная шкала от 15% до 20%."),
            ("3. Минимальный порог", "При результате ниже 80% бонус не выплачивается. Возвраты клиента до даты расчёта уменьшают зачтённую выручку."),
        ],
    )
    create_pdf(
        "additional_agreement_remote_engineering.pdf",
        "RW-ENG-2025",
        "Режим удалённой работы инженерного отдела",
        "1 августа 2025 года; действует до 31 декабря 2026 года.",
        "Инженеры разработки, тестирования и платформенной команды продуктового блока.",
        "Имеет приоритет над базовым лимитом удалённой работы; обязательные очные мероприятия сохраняются.",
        [
            ("1. Специальный режим", "Допускается до четырёх удалённых рабочих дней в неделю. Среда является общим офисным днём для инженерного отдела."),
            ("2. Дежурства", "Во время производственного дежурства место работы определяется графиком дежурств. Аварийный выезд в дата-центр не считается изменением постоянного режима."),
            ("3. Безопасность", "Доступ к исходному коду выполняется только с управляемого устройства через корпоративную VPN и многофакторную аутентификацию."),
        ],
    )
    create_pdf(
        "additional_agreement_bonus_product.pdf",
        "BONUS-PRODUCT-Q3-2025",
        "Квартальный бонус продуктового отдела",
        "1 июля 2025 года; применяется к результатам III квартала 2025 года.",
        "Менеджеры продукта и продуктовые аналитики, включённые в утверждённые команды релиза.",
        "Имеет приоритет над базовым целевым бонусом 10% при расчёте за III квартал 2025 года.",
        [
            ("1. Целевой бонус", "Целевой бонус составляет 12% квартального оклада при своевременном выпуске утверждённого релиза и выполнении индивидуальных целей на 100%."),
            ("2. Условия выплаты", "Задержка релиза более чем на десять календарных дней по причинам в зоне контроля команды снижает коэффициент релиза до 0,7."),
            ("3. Исключения", "Перенос срока, письменно утверждённый продуктовым директором до контрольной даты, не считается задержкой."),
        ],
    )
    create_pdf(
        "additional_agreement_travel_north_regions.pdf",
        "TR-NORTH-2025",
        "Лимит гостиницы для северных регионов",
        "1 сентября 2025 года; действует до 31 марта 2027 года.",
        "Все штатные сотрудники в командировках в Норильск, Мурманск, Новый Уренгой и Якутск, кроме руководителей высшего звена.",
        "Имеет приоритет над базовыми гостиничными лимитами по России; для руководителей высшего звена действует отдельное соглашение.",
        [
            ("1. Специальный лимит", "Максимальная стоимость стандартного гостиничного номера составляет 14 000 рублей за ночь, включая завтрак и обязательные местные сборы."),
            ("2. Подтверждение", "К авансовому отчёту прикладываются счёт и кассовый чек. Превышение лимита требует предварительного письменного согласования финансового контролёра."),
            ("3. Остальные условия", "Суточные, транспортные расходы и сроки подачи авансового отчёта регулируются базовой политикой командировок."),
        ],
    )


if __name__ == "__main__":
    main()
