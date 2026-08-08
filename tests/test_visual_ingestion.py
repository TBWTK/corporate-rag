from __future__ import annotations

import json
import shutil
from pathlib import Path

import pytest
from docx import Document as WordDocument
from PIL import Image
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from sqlalchemy import select
from sqlalchemy.orm import Session, sessionmaker

from rag_app.config import Settings
from rag_app.db.models import Chunk, Document, DocumentStatus, KnowledgeSpace
from rag_app.ingestion.vision import build_page_prompt, parse_visual_page
from rag_app.ingestion.visual import (
    RenderedPage,
    document_has_visuals,
    page_image_path,
    page_number_from_location,
    render_document_pages,
)
from rag_app.providers.base import Completion
from rag_app.providers.fake import FakeProvider
from rag_app.services.ingestion import IngestionWorker, UploadService


def _make_image(path: Path) -> None:
    Image.new("RGB", (240, 100), color=(228, 242, 235)).save(path)


def _make_visual_pdf(path: Path, image_path: Path, *, pages: int = 1) -> None:
    pdf = canvas.Canvas(str(path))
    for number in range(1, pages + 1):
        pdf.drawString(72, 760, f"Step {number}: open Settings")
        pdf.drawImage(ImageReader(str(image_path)), 72, 600, width=240, height=100)
        pdf.showPage()
    pdf.save()


def _make_visual_docx(path: Path, image_path: Path) -> None:
    document = WordDocument()
    document.add_heading("Настройка почты", level=1)
    document.add_paragraph("Откройте раздел Параметры.")
    document.add_picture(str(image_path))
    document.save(path)


def test_detects_embedded_visuals_in_pdf_and_docx(tmp_path: Path) -> None:
    image = tmp_path / "screen.png"
    pdf = tmp_path / "guide.pdf"
    docx = tmp_path / "guide.docx"
    _make_image(image)
    _make_visual_pdf(pdf, image)
    _make_visual_docx(docx, image)

    assert document_has_visuals(pdf) is True
    assert document_has_visuals(docx) is True

    plain_docx = tmp_path / "plain.docx"
    WordDocument().save(plain_docx)
    assert document_has_visuals(plain_docx) is False


@pytest.mark.skipif(
    not shutil.which("pdftoppm") or not shutil.which("soffice"),
    reason="Poppler and LibreOffice are required for page render",
)
def test_renders_pdf_and_docx_into_ordered_page_images(tmp_path: Path) -> None:
    image = tmp_path / "screen.png"
    pdf = tmp_path / "guide.pdf"
    docx = tmp_path / "guide.docx"
    _make_image(image)
    _make_visual_pdf(pdf, image, pages=2)
    _make_visual_docx(docx, image)

    pdf_pages = render_document_pages(pdf, tmp_path / "pdf-pages", dpi=96)
    docx_pages = render_document_pages(docx, tmp_path / "docx-pages", dpi=96)

    assert [page.number for page in pdf_pages] == [1, 2]
    assert [page.number for page in docx_pages] == [1]
    assert all(page.path.stat().st_size > 0 for page in (*pdf_pages, *docx_pages))
    assert page_image_path(pdf, 1).name == "page-1.png"
    assert pdf_pages[0].path.name == "page-1.png"


def test_normalizes_visual_json_into_page_overview_and_steps() -> None:
    raw = json.dumps(
        {
            "page_title": "Настройка учётной записи",
            "page_summary": "Добавление почтового сервера",
            "visible_text": ["mail.example.ru", "Порт 443"],
            "steps": [
                {
                    "step_number": "3",
                    "action": "Нажмите «Добавить учётную запись»",
                    "ui_target": "Кнопка «Добавить»",
                    "value": "mail.example.ru",
                    "expected_result": "Откроется форма входа",
                    "warning": "Не используйте личный пароль",
                }
            ],
        },
        ensure_ascii=False,
    )

    units = parse_visual_page(raw, page_number=4)

    assert [unit.location for unit in units] == ["стр. 4 · шаг 3"]
    assert "mail.example.ru" in units[0].text
    assert "Кнопка «Добавить»" in units[0].text
    assert "Не используйте" in units[0].text
    assert page_number_from_location(units[0].location) == 4
    assert "Текстовый слой" in build_page_prompt("guide.pdf", 4, "native content")


def test_visual_parser_keeps_non_json_response_as_page_fallback() -> None:
    units = parse_visual_page("На странице показана кнопка Настройки", page_number=2)

    assert len(units) == 1
    assert units[0].location == "стр. 2 · визуальный разбор"


@pytest.mark.skipif(
    not shutil.which("pdftoppm"), reason="Poppler is required for visual worker test"
)
def test_worker_indexes_visual_steps_and_keeps_page_image(
    tmp_path: Path, sqlite_factory: sessionmaker[Session]
) -> None:
    settings = Settings(
        _env_file=None,
        llm_provider="fake",
        data_dir=tmp_path / "uploads",
        visual_page_dpi=96,
    )
    with sqlite_factory() as session:
        space = KnowledgeSpace(name="Инструкции")
        session.add(space)
        session.commit()
        session.refresh(space)

    image = tmp_path / "screen.png"
    source = tmp_path / "mail.pdf"
    _make_image(image)
    _make_visual_pdf(source, image)
    queued = UploadService(settings, sqlite_factory).queue(
        space_id=space.id,
        filename="mail.pdf",
        media_type="application/pdf",
        payload=source.read_bytes(),
    )

    class VisionFake(FakeProvider):
        def __init__(self) -> None:
            super().__init__()
            self.images: list[Path] = []

        def analyze_image(self, image_path: Path, *, prompt: str) -> Completion:
            assert "Страница: 1" in prompt
            self.images.append(image_path)
            return Completion(
                text=json.dumps(
                    {
                        "page_title": "Почта",
                        "page_summary": "Настройка сервера",
                        "visible_text": ["mail.example.ru"],
                        "steps": [
                            {
                                "step_number": "1",
                                "action": "Введите адрес сервера",
                                "ui_target": "Поле «Сервер»",
                                "value": "mail.example.ru",
                                "expected_result": "Подключение установлено",
                                "warning": "",
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                model="fake-vision",
            )

    provider = VisionFake()
    IngestionWorker(settings, sqlite_factory, provider).process(queued.id)

    with sqlite_factory() as session:
        document = session.get(Document, queued.id)
        chunks = list(
            session.scalars(
                select(Chunk).where(Chunk.document_id == queued.id).order_by(Chunk.chunk_index)
            )
        )
    assert document is not None and document.status == DocumentStatus.READY
    assert any(chunk.location == "стр. 1 · шаг 1" for chunk in chunks)
    assert any("mail.example.ru" in chunk.content for chunk in chunks)
    assert len(provider.images) == 1
    assert provider.images[0].is_file()


def test_worker_uses_text_fallback_when_provider_has_no_vision(
    tmp_path: Path,
    sqlite_factory: sessionmaker[Session],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    settings = Settings(
        _env_file=None,
        llm_provider="fake",
        data_dir=tmp_path / "uploads",
    )
    with sqlite_factory() as session:
        space = KnowledgeSpace(name="Fallback")
        session.add(space)
        session.commit()
        session.refresh(space)
    image = tmp_path / "screen.png"
    source = tmp_path / "fallback.pdf"
    _make_image(image)
    _make_visual_pdf(source, image)
    queued = UploadService(settings, sqlite_factory).queue(
        space_id=space.id,
        filename="fallback.pdf",
        media_type="application/pdf",
        payload=source.read_bytes(),
    )

    def fail_render(*_args: object, **_kwargs: object) -> tuple[RenderedPage, ...]:
        raise AssertionError("render should not run without vision provider")

    monkeypatch.setattr("rag_app.services.ingestion.render_document_pages", fail_render)
    IngestionWorker(settings, sqlite_factory, FakeProvider()).process(queued.id)

    with sqlite_factory() as session:
        document = session.get(Document, queued.id)
    assert document is not None and document.status == DocumentStatus.READY
